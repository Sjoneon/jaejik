# ============================================
# 업무 일정 관리 시스템 - 문서 파싱 서비스
# 위치: C:\Users\user\Desktop\인공지능산업협회AI\services\document_parser.py
# ============================================

import os
from typing import Optional, Tuple


class DocumentParser:
    """문서 파싱 서비스 - HWP, DOCX, PDF, Excel 지원"""
    
    SUPPORTED_EXTENSIONS = {'hwp', 'hwpx', 'docx', 'doc', 'pdf', 'xlsx', 'xls', 'csv'}
    
    @classmethod
    def parse(cls, filepath: str) -> Tuple[bool, str, Optional[str]]:
        """
        문서 파싱 메인 함수
        
        Args:
            filepath: 파일 경로
            
        Returns:
            Tuple[성공여부, 메시지, 추출된텍스트]
        """
        if not filepath or not os.path.exists(filepath):
            return False, "파일을 찾을 수 없습니다.", None
        
        # 확장자 추출
        ext = cls._get_extension(filepath)
        
        if ext not in cls.SUPPORTED_EXTENSIONS:
            return False, f"지원하지 않는 파일 형식입니다: {ext}", None
        
        try:
            # 확장자별 파싱
            if ext in ('hwp', 'hwpx'):
                return cls._parse_hwp(filepath)
            elif ext in ('docx', 'doc'):
                return cls._parse_docx(filepath)
            elif ext == 'pdf':
                return cls._parse_pdf(filepath)
            elif ext in ('xlsx', 'xls'):
                return cls._parse_excel(filepath)
            elif ext == 'csv':
                return cls._parse_csv(filepath)
            else:
                return False, "알 수 없는 파일 형식입니다.", None
                
        except Exception as e:
            return False, f"파일 파싱 중 오류 발생: {str(e)}", None
    
    @staticmethod
    def _get_extension(filepath: str) -> str:
        """파일 확장자 추출"""
        if not filepath:
            return ""
        
        # 방법 1: os.path.splitext 사용
        _, ext = os.path.splitext(filepath)
        if ext:
            return ext.lower().lstrip('.')
        
        # 방법 2: 파일명에서 직접 추출 (백업)
        filename = os.path.basename(filepath)
        if '.' in filename:
            return filename.rsplit('.', 1)[-1].lower()
        
        return ""
    
    @classmethod
    def _parse_excel(cls, filepath: str) -> Tuple[bool, str, Optional[str]]:
        """Excel 파일 파싱 (xlsx, xls)"""
        try:
            import openpyxl
        except ImportError:
            return False, "openpyxl 라이브러리가 설치되지 않았습니다. pip install openpyxl", None
        
        try:
            wb = openpyxl.load_workbook(filepath, data_only=True)
            
            extracted_text = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                extracted_text.append(f"[시트: {sheet_name}]")
                
                for row in sheet.iter_rows():
                    row_values = []
                    for cell in row:
                        if cell.value is not None:
                            row_values.append(str(cell.value))
                    
                    if row_values:
                        extracted_text.append(' | '.join(row_values))
                
                extracted_text.append("")  # 시트 구분
            
            wb.close()
            
            full_text = '\n'.join(extracted_text)
            
            if full_text.strip():
                return True, "Excel 파일 파싱 성공", full_text.strip()
            else:
                return True, "Excel 파일을 열었으나 데이터가 없습니다.", ""
                
        except Exception as e:
            return False, f"Excel 파싱 오류: {str(e)}", None
    
    @classmethod
    def _parse_csv(cls, filepath: str) -> Tuple[bool, str, Optional[str]]:
        """CSV 파일 파싱"""
        try:
            extracted_text = []
            
            # 인코딩 시도 순서
            encodings = ['utf-8', 'cp949', 'euc-kr', 'utf-8-sig']
            
            content = None
            for encoding in encodings:
                try:
                    with open(filepath, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                return False, "CSV 파일 인코딩을 인식할 수 없습니다.", None
            
            # 줄 단위로 처리
            for line in content.split('\n'):
                line = line.strip()
                if line:
                    # 쉼표를 | 로 변환해서 가독성 향상
                    extracted_text.append(line.replace(',', ' | '))
            
            full_text = '\n'.join(extracted_text)
            
            if full_text.strip():
                return True, "CSV 파일 파싱 성공", full_text.strip()
            else:
                return True, "CSV 파일을 열었으나 데이터가 없습니다.", ""
                
        except Exception as e:
            return False, f"CSV 파싱 오류: {str(e)}", None
    
    @classmethod
    def _parse_hwp(cls, filepath: str) -> Tuple[bool, str, Optional[str]]:
        """
        HWP 파일 파싱
        
        HWP는 복잡한 바이너리 포맷이므로 olefile을 사용해 기본 텍스트 추출
        """
        try:
            import olefile
        except ImportError:
            return False, "olefile 라이브러리가 설치되지 않았습니다. pip install olefile", None
        
        try:
            # HWP 파일 열기
            ole = olefile.OleFileIO(filepath)
            
            # 본문 텍스트 스트림 찾기
            extracted_text = ""
            
            # HWP의 PrvText (미리보기 텍스트) 스트림에서 텍스트 추출
            if ole.exists('PrvText'):
                prvtext_stream = ole.openstream('PrvText')
                prvtext_data = prvtext_stream.read()
                
                # UTF-16 LE로 디코딩 (HWP 기본 인코딩)
                try:
                    extracted_text = prvtext_data.decode('utf-16-le', errors='ignore')
                except UnicodeDecodeError:
                    extracted_text = prvtext_data.decode('cp949', errors='ignore')
            
            # PrvText가 없으면 다른 방법 시도
            if not extracted_text.strip():
                # BodyText/Section0 등에서 추출 시도
                for entry in ole.listdir():
                    entry_name = '/'.join(entry)
                    if 'bodytext' in entry_name.lower() or 'section' in entry_name.lower():
                        try:
                            stream = ole.openstream(entry)
                            data = stream.read()
                            # 바이너리에서 한글 텍스트 추출 시도
                            text = cls._extract_korean_text(data)
                            if text:
                                extracted_text += text + "\n"
                        except Exception:
                            continue
            
            ole.close()
            
            if extracted_text.strip():
                return True, "HWP 파일 파싱 성공", extracted_text.strip()
            else:
                return True, "HWP 파일을 열었으나 텍스트를 추출하지 못했습니다.", ""
                
        except Exception as e:
            return False, f"HWP 파싱 오류: {str(e)}", None
    
    @staticmethod
    def _extract_korean_text(data: bytes) -> str:
        """바이너리 데이터에서 한글 텍스트 추출 시도"""
        result = []
        
        # UTF-16 LE로 디코딩 시도
        try:
            text = data.decode('utf-16-le', errors='ignore')
            # 출력 가능한 문자만 필터링
            filtered = ''.join(c for c in text if c.isprintable() or c in '\n\t ')
            if len(filtered) > 10:
                return filtered
        except Exception:
            pass
        
        # CP949로 디코딩 시도
        try:
            text = data.decode('cp949', errors='ignore')
            filtered = ''.join(c for c in text if c.isprintable() or c in '\n\t ')
            if len(filtered) > 10:
                return filtered
        except Exception:
            pass
        
        return ""
    
    @classmethod
    def _parse_docx(cls, filepath: str) -> Tuple[bool, str, Optional[str]]:
        """DOCX 파일 파싱"""
        try:
            from docx import Document
        except ImportError:
            return False, "python-docx 라이브러리가 설치되지 않았습니다. pip install python-docx", None
        
        try:
            doc = Document(filepath)
            
            paragraphs = []
            
            # 모든 단락 추출
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # 표에서도 텍스트 추출
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text and text not in paragraphs:
                            paragraphs.append(text)
            
            extracted_text = '\n'.join(paragraphs)
            
            if extracted_text:
                return True, "DOCX 파일 파싱 성공", extracted_text
            else:
                return True, "DOCX 파일을 열었으나 텍스트가 없습니다.", ""
                
        except Exception as e:
            return False, f"DOCX 파싱 오류: {str(e)}", None
    
    @classmethod
    def _parse_pdf(cls, filepath: str) -> Tuple[bool, str, Optional[str]]:
        """PDF 파일 파싱"""
        try:
            import pdfplumber
        except ImportError:
            return False, "pdfplumber 라이브러리가 설치되지 않았습니다. pip install pdfplumber", None
        
        try:
            extracted_text = []
            
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        extracted_text.append(text)
            
            full_text = '\n'.join(extracted_text)
            
            if full_text.strip():
                return True, "PDF 파일 파싱 성공", full_text.strip()
            else:
                return True, "PDF 파일을 열었으나 텍스트가 없습니다.", ""
                
        except Exception as e:
            return False, f"PDF 파싱 오류: {str(e)}", None
    
    @classmethod
    def is_supported(cls, filename: str) -> bool:
        """지원하는 파일 형식인지 확인"""
        ext = cls._get_extension(filename)
        return ext in cls.SUPPORTED_EXTENSIONS
